"""
material-analysis / extract.py

원본 자료(docx/pdf/xlsx/csv/이미지)에서 텍스트·표·이미지를 추출해
material_analysis.json 스키마의 원재료(sources)를 생성한다.

## 계층 보존 원칙 (2026-08-19 개정)
기존 버전은 문단/표/이미지를 각각 flat한 `text_blocks[]`/`tables[]`/`images[]`
배열로만 추출해, 원본에 존재하던 "무엇이 어느 상위 주제 아래에 속하는가"라는
구조 정보가 이 시점에 이미 사라졌다. 이번 개정은 **원본에서 실제로 확인 가능한
구조 신호가 있을 때만** 그 신호를 그대로 보존해 `content_groups[] → subtopics[]`
계층으로 추출한다. 신호가 없으면 계층을 임의로 만들지 않고 `structure_signal:
"none_detected"`로 표시한 뒤 flat 구조로 남긴다(추정 금지).

형식별로 사용하는 구조 신호는 다르다 — 특정 신호(예: Word Heading 스타일)에만
의존하지 않는다:
- docx: 문단 스타일 이름의 "Heading N"/"제목 N" 패턴(다국어 템플릿 대응),
  문서 본문(`document.xml`)의 실제 XML 순서(문단·표·이미지가 섞여 있는 원래
  순서 그대로 순회 — 이전 버전은 `doc.tables`/`doc.part.rels`를 별도 컬렉션으로
  순회해 실제 문서 순서와 다를 수 있었다)
- pdf: 페이지 경계(1차 신호, 항상 사용 가능) + 페이지 내 폰트 크기 이상치
  휴리스틱(2차 신호, "font_size_heuristic"로 명시 — 확정 신호가 아님)
- xlsx: 시트 경계
- csv/이미지 단일 파일: 구조 신호 없음(`none_detected`) — 강제로 계층을 만들지
  않는다

## 원칙(설계서 1.6 준수, 기존 그대로 유지)
- 수치는 원본 그대로 추출한다. 재계산·추정하지 않는다.
- 스캔 PDF(텍스트 레이어 없음) 등 자동 판독이 불가능한 경우 OCR을 임의로
  시도하지 않고 needs_manual_review 플래그만 남긴다.
- 임베드 차트는 원본 수치를 추출하지 않고 "원본 활용 필요"로만 표시한다.
- 이미지는 파일 경로와 (확인 가능한 경우) 소속 Content Group/Subtopic까지만
  기계적으로 기록한다 — 이미지가 "무엇을 뒷받침하는지"의 의미 판단, 그리고
  라벨의 신뢰도(Uncertain 여부) 판단은 이 스크립트의 책임이 아니라
  material-analysis SKILL.md 2단계(LLM)의 책임이다.

사용법:
    python extract.py --input <파일 또는 폴더> [<파일2> ...] --output <json 경로>
"""
from __future__ import annotations

import argparse
import json
import re
from collections import Counter
from pathlib import Path
from typing import Any

SUPPORTED_EXTS = {".docx", ".pdf", ".xlsx", ".csv", ".png", ".jpg", ".jpeg", ".gif", ".bmp"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".bmp"}

# Word 스타일 이름에서 헤딩 레벨을 뽑아내는 패턴. 특정 로캘의 한 가지 이름
# ("Heading 1")에만 의존하지 않도록 영문/국문 두 표기를 모두 인식한다.
_HEADING_STYLE_RE = re.compile(r"^(Heading|제목)\s*([0-9]+)$", re.IGNORECASE)
_TITLE_STYLE_NAMES = {"title", "표제", "부제"}


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _heading_level(style_name: str | None) -> int | None:
    """스타일 이름에서 헤딩 레벨을 판단한다. 헤딩이 아니면 None."""
    if not style_name:
        return None
    if style_name.strip().lower() in _TITLE_STYLE_NAMES:
        return 0
    m = _HEADING_STYLE_RE.match(style_name.strip())
    if m:
        return int(m.group(2))
    return None


def _iter_block_items(document):
    """문서 본문(body)을 실제 XML 순서 그대로(문단·표가 섞인 순서) 순회한다.

    python-docx의 `document.paragraphs`/`document.tables`는 각각 별도
    컬렉션이라 원래 문서에서 문단과 표가 번갈아 나오는 순서를 재구성할 수
    없다 — 이 순서(어떤 문단 바로 다음에 어떤 표/이미지가 오는가)가 바로
    "이 표/이미지가 어느 제목 아래에 있었는가"를 판단하는 유일한 근거이므로
    body를 직접 순회해 복원한다.
    """
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _paragraph_image_rids(paragraph) -> list[str]:
    """문단 안에 인라인으로 삽입된 이미지의 relationship id를 문서 순서대로 반환."""
    from docx.oxml.ns import qn

    rids = []
    for blip in paragraph._p.findall(".//" + qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if rid:
            rids.append(rid)
    return rids


def _cell_image_rids(cell) -> list[str]:
    """표 셀 안에 삽입된 이미지의 relationship id를 순서대로 반환.

    변환 도구(PPT→docx 등)로 만들어진 문서는 이미지 한 장을 담기 위해 셀 1개짜리
    표를 쓰는 경우가 흔하다 — 이 경우 표를 순회하지 않으면 이미지 자체가
    "표"로 오분류되어 유실된다.
    """
    from docx.oxml.ns import qn

    rids = []
    for blip in cell._tc.findall(".//" + qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if rid:
            rids.append(rid)
    return rids


def _textbox_texts(element, outer_text: str = "") -> list[str]:
    """Return text from DrawingML text boxes omitted by python-docx text APIs."""
    wps_txbx = "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx"
    w_p = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
    w_t = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"

    texts: list[str] = []
    for textbox in element.iter(wps_txbx):
        # Preserve paragraph order inside each text box. Only wps:txbx is
        # traversed, so ordinary body/table text stays on its existing path.
        for paragraph in textbox.iter(w_p):
            text = "".join(node.text or "" for node in paragraph.iter(w_t)).strip()
            # Prevent duplicate output if a future python-docx release exposes
            # nested text-box runs through Paragraph.text or _Cell.text.
            if text and text not in outer_text:
                texts.append(text)
    return texts


def _save_image_rel(rid: str, rels, image_dir: Path, stem: str, seq: int) -> str | None:
    """relationship id가 가리키는 이미지를 저장하고 경로를 반환한다(이미지가 아니면 None)."""
    if rid not in rels:
        return None
    rel = rels[rid]
    if "image" not in rel.reltype:
        return None
    image_part = rel.target_part
    ext = image_part.content_type.split("/")[-1]
    out_path = image_dir / f"{stem}_img{seq}.{ext}"
    out_path.write_bytes(image_part.blob)
    return str(out_path)


def _new_group(index: int, title: str, signal_detail: str) -> dict[str, Any]:
    return {
        "group_index": index,
        "title": title,
        "signal_detail": signal_detail,
        "subtopics": [],
        "direct_text_blocks": [],
        "direct_tables": [],
        "direct_images": [],
    }


def _new_subtopic(index: int, title: str, signal_detail: str) -> dict[str, Any]:
    return {
        "subtopic_index": index,
        "title": title,
        "signal_detail": signal_detail,
        "text_blocks": [],
        "tables": [],
        "images": [],
    }


def extract_docx(path: Path, image_dir: Path) -> dict[str, Any]:
    from docx import Document

    doc = Document(str(path))
    rels = doc.part.rels

    content_groups: list[dict[str, Any]] = []
    unstructured_text_blocks: list[dict[str, Any]] = []
    unstructured_tables: list[dict[str, Any]] = []
    unstructured_images: list[dict[str, Any]] = []

    current_group: dict[str, Any] | None = None
    current_subtopic: dict[str, Any] | None = None
    group_level: int | None = None  # 최초로 발견된(가장 얕은) 헤딩 레벨 = Content Group 경계
    subtopic_level: int | None = None  # 그 다음 깊이 = Subtopic 경계
    order = 0
    image_seq = 0
    heading_found = False

    def _target(kind: str):
        """현재 위치(group/subtopic 유무)에 맞는 저장 대상 리스트를 반환."""
        if current_subtopic is not None:
            return current_subtopic[kind]
        if current_group is not None:
            return current_group["direct_" + kind]
        return {"text_blocks": unstructured_text_blocks, "tables": unstructured_tables,
                "images": unstructured_images}[kind]

    for block in _iter_block_items(doc):
        is_table = block.__class__.__name__ == "Table"

        if not is_table:
            para = block
            text = para.text.strip()
            style_name = para.style.name if para.style else None
            level = _heading_level(style_name)

            if level is not None and level > 0:
                heading_found = True
                if group_level is None:
                    group_level = level
                if level == group_level:
                    order += 1
                    current_group = _new_group(len(content_groups) + 1, text, f"{style_name}")
                    content_groups.append(current_group)
                    current_subtopic = None
                elif subtopic_level is None or level == subtopic_level:
                    subtopic_level = level
                    order += 1
                    if current_group is None:
                        # 상위 그룹 헤딩 없이 subtopic 레벨 헤딩만 먼저 나오는 예외적인 경우 —
                        # 이 subtopic을 담을 임시 그룹을 만든다(원본 구조를 잃지 않기 위함).
                        current_group = _new_group(len(content_groups) + 1, "(제목 없음)", "묵시적 그룹")
                        content_groups.append(current_group)
                    current_subtopic = _new_subtopic(len(current_group["subtopics"]) + 1, text, f"{style_name}")
                    current_group["subtopics"].append(current_subtopic)
                else:
                    # group/subtopic보다 더 깊은 레벨의 헤딩 — 별도 3번째 계층을 만들지 않고
                    # 현재 subtopic의 본문 텍스트에 "[하위 제목] ..." 형태로 유실 없이 남긴다.
                    order += 1
                    _target("text_blocks").append({"text": f"[{text}]", "order": order, "role": "sub_heading"})
                continue

            if level == 0:
                # 문서 Title — 어느 Content Group에도 속하지 않는 문서 전체 제목이므로
                # 별도 처리(계층에 넣지 않고 최상위에 기록, main()에서 사용).
                order += 1
                continue

            if text:
                order += 1
                _target("text_blocks").append({"text": text, "order": order, "style": style_name})

            # Paragraph.text omits text nested in DrawingML text boxes. Keep
            # it as ordinary evidence, never as a heading.
            for textbox_text in _textbox_texts(para._p, text):
                order += 1
                _target("text_blocks").append({
                    "text": textbox_text,
                    "order": order,
                    "style": None,
                    "source_type": "text_box",
                })

            # 문단 안의 인라인 이미지를 문서 순서 그대로 추출 — 이 순간의
            # current_group/current_subtopic이 바로 이 이미지의 소속 근거(위치 신호)다.
            for rid in _paragraph_image_rids(para):
                image_seq += 1
                out = _save_image_rel(rid, rels, image_dir, path.stem, image_seq)
                if out is None:
                    continue
                order += 1
                caption = text if text else None  # 같은 문단에 캡션 성격 텍스트가 있으면 함께 기록
                _target("images").append({
                    "path": out,
                    "order": order,
                    "adjacent_text": caption,
                })
        else:
            table = block
            rows_cells = [[cell for cell in row.cells] for row in table.rows]
            rows_text = [[cell.text for cell in row_cells] for row_cells in rows_cells]

            # 변환 도구가 이미지 한 장을 셀 1개짜리 표로 감싼 "이미지 컨테이너" 표인지
            # 판단한다 — 모든 셀의 텍스트가 비어 있으면(=원본 자료에 실제 표 데이터가
            # 없으면) 표로 취급하지 않고, 셀 안의 이미지만 순서대로 꺼낸다.
            is_pure_image_container = bool(rows_cells) and all(
                not cell_text.strip() for row_text in rows_text for cell_text in row_text
            )

            if is_pure_image_container:
                for row_cells in rows_cells:
                    for cell in row_cells:
                        for rid in _cell_image_rids(cell):
                            image_seq += 1
                            out = _save_image_rel(rid, rels, image_dir, path.stem, image_seq)
                            if out is None:
                                continue
                            order += 1
                            _target("images").append({"path": out, "order": order, "adjacent_text": None})
            else:
                headers = rows_text[0] if rows_text else []
                body_rows = rows_text[1:] if len(rows_text) > 1 else []
                order += 1
                _target("tables").append({
                    "headers": headers,
                    "rows": body_rows,
                    "row_count": len(body_rows),
                    "needs_appendix": len(body_rows) >= 30,
                    "order": order,
                })
                # 데이터 표 안에 사진이 함께 삽입된(드문) 경우도 유실 없이 함께 추출한다.
                for row_cells in rows_cells:
                    for cell in row_cells:
                        for rid in _cell_image_rids(cell):
                            image_seq += 1
                            out = _save_image_rel(rid, rels, image_dir, path.stem, image_seq)
                            if out is None:
                                continue
                            order += 1
                            _target("images").append({"path": out, "order": order, "adjacent_text": None})

            # _Cell.text also omits wps:txbx. Run this outside the table type
            # branch so text boxes in image-container tables are retained too.
            for row_cells in rows_cells:
                for cell in row_cells:
                    for textbox_text in _textbox_texts(cell._tc, cell.text):
                        order += 1
                        _target("text_blocks").append({
                            "text": textbox_text,
                            "order": order,
                            "style": None,
                            "source_type": "text_box",
                        })

    return {
        "file": str(path),
        "type": "docx",
        "structure_signal": "heading_style" if heading_found else "none_detected",
        "content_groups": content_groups,
        "unstructured_text_blocks": unstructured_text_blocks,
        "unstructured_tables": unstructured_tables,
        "unstructured_images": unstructured_images,
        "needs_manual_review": False,
    }


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _pdf_page_heading_candidates(page, body_size: float) -> list[dict[str, Any]]:
    """페이지 내 줄 단위로 폰트 크기를 계산해, 본문보다 눈에 띄게 큰 줄을
    '제목 후보'로 표시한다. 확정 신호가 아니라 휴리스틱이므로 상위 스키마에
    signal_detail: "font_size_heuristic"로 명시해 하위 판단이 이를 확정된
    계층으로 오인하지 않도록 한다.
    """
    words = page.extract_words(extra_attrs=["size"]) or []
    if not words:
        return []
    lines: dict[int, list[dict]] = {}
    for w in words:
        key = round(w["top"])
        lines.setdefault(key, []).append(w)

    candidates = []
    threshold = body_size * 1.15 if body_size else None
    for top, ws in sorted(lines.items()):
        avg_size = sum(w.get("size", 0) for w in ws) / len(ws)
        if threshold and avg_size >= threshold:
            text = " ".join(w["text"] for w in sorted(ws, key=lambda w: w["x0"]))
            candidates.append({"text": text, "top": top, "size": round(avg_size, 1)})
    return candidates


def extract_pdf(path: Path, image_dir: Path) -> dict[str, Any]:
    import pdfplumber

    content_groups: list[dict[str, Any]] = []
    needs_manual_review = False

    with pdfplumber.open(str(path)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text() or ""
            if not page_text.strip():
                needs_manual_review = True

            all_chars = page.chars or []
            body_size = 0.0
            if all_chars:
                size_counts = Counter(round(c.get("size", 0), 1) for c in all_chars)
                body_size = size_counts.most_common(1)[0][0]

            heading_candidates = _pdf_page_heading_candidates(page, body_size) if all_chars else []
            group_title = heading_candidates[0]["text"] if heading_candidates else f"Page {page_num}"
            signal = "font_size_heuristic" if heading_candidates else "page_boundary"

            group = _new_group(page_num, group_title, signal)
            content_groups.append(group)

            if page_text.strip():
                group["direct_text_blocks"].append({"text": page_text, "order": 1})

            for t_idx, table in enumerate(page.extract_tables(), start=1):
                if not table:
                    continue
                headers, *body_rows = table
                group["direct_tables"].append({
                    "headers": headers,
                    "rows": body_rows,
                    "row_count": len(body_rows),
                    "needs_appendix": len(body_rows) >= 30,
                    "order": t_idx,
                })

            for img_idx, img in enumerate(page.images, start=1):
                try:
                    cropped = page.within_bbox((img["x0"], img["top"], img["x1"], img["bottom"]))
                    im = cropped.to_image(resolution=200)
                    out_path = image_dir / f"{path.stem}_p{page_num}_img{img_idx}.png"
                    im.save(str(out_path))
                    group["direct_images"].append({"path": str(out_path), "order": img_idx})
                except Exception:
                    # 페이지 이미지 추출 실패는 치명적이지 않음 — 건너뛰고 계속 진행
                    continue

    return {
        "file": str(path),
        "type": "pdf",
        "structure_signal": "page_boundary",
        "content_groups": content_groups,
        "unstructured_text_blocks": [],
        "unstructured_tables": [],
        "unstructured_images": [],
        "needs_manual_review": needs_manual_review,
    }


# ---------------------------------------------------------------------------
# XLSX / CSV / 이미지
# ---------------------------------------------------------------------------

def extract_xlsx(path: Path, image_dir: Path) -> dict[str, Any]:
    import openpyxl

    wb = openpyxl.load_workbook(str(path), data_only=True)
    content_groups: list[dict[str, Any]] = []
    charts_detected = []

    for s_idx, sheet in enumerate(wb.worksheets, start=1):
        group = _new_group(s_idx, sheet.title, "sheet_boundary")
        content_groups.append(group)

        rows = [[cell for cell in row] for row in sheet.iter_rows(values_only=True)]
        rows = [r for r in rows if any(v is not None for v in r)]
        if rows:
            headers = list(rows[0])
            body = [list(r) for r in rows[1:]]
            group["direct_tables"].append({
                "headers": headers,
                "rows": body,
                "row_count": len(body),
                "needs_appendix": len(body) >= 30,
                "order": 1,
            })

        if getattr(sheet, "_charts", None):
            for _ in sheet._charts:
                charts_detected.append({
                    "location": f"{path.name}!{sheet.title}",
                    "note": "차트 감지됨 — 원본 수치 추출 불가. 원본 파일 활용 또는 사용자에게 원본 데이터 확인 필요",
                })

    return {
        "file": str(path),
        "type": "xlsx",
        "structure_signal": "sheet_boundary",
        "content_groups": content_groups,
        "unstructured_text_blocks": [],
        "unstructured_tables": [],
        "unstructured_images": [],
        "charts_detected": charts_detected,
        "needs_manual_review": False,
    }


def extract_csv(path: Path, image_dir: Path) -> dict[str, Any]:
    import csv

    with open(path, newline="", encoding="utf-8-sig") as f:
        reader = list(csv.reader(f))

    headers = reader[0] if reader else []
    body = reader[1:] if len(reader) > 1 else []

    group = _new_group(1, path.stem, "none_detected")
    group["direct_tables"].append({
        "headers": headers,
        "rows": body,
        "row_count": len(body),
        "needs_appendix": len(body) >= 30,
        "order": 1,
    })

    return {
        "file": str(path),
        "type": "csv",
        "structure_signal": "none_detected",
        "content_groups": [group],
        "unstructured_text_blocks": [],
        "unstructured_tables": [],
        "unstructured_images": [],
        "needs_manual_review": False,
    }


def extract_image(path: Path, image_dir: Path) -> dict[str, Any]:
    # 이미지 자체의 내용 판단(무엇을 담고 있는지)은 LLM이 이미지를 직접 보고 수행한다.
    # 단일 이미지 파일은 원본 구조 신호가 전혀 없으므로 계층을 만들지 않는다.
    return {
        "file": str(path),
        "type": "image",
        "structure_signal": "none_detected",
        "content_groups": [],
        "unstructured_text_blocks": [],
        "unstructured_tables": [],
        "unstructured_images": [{"path": str(path), "order": 1}],
        "needs_manual_review": False,
    }


EXTRACTORS = {
    ".docx": extract_docx,
    ".pdf": extract_pdf,
    ".xlsx": extract_xlsx,
    ".csv": extract_csv,
}


def collect_files(inputs: list[str]) -> list[Path]:
    files: list[Path] = []
    for raw in inputs:
        p = Path(raw)
        if p.is_dir():
            files.extend(sorted(x for x in p.rglob("*") if x.suffix.lower() in SUPPORTED_EXTS))
        elif p.is_file():
            files.append(p)
        else:
            raise FileNotFoundError(f"입력 경로를 찾을 수 없습니다: {raw}")
    return files


def build_escalations_and_confirmations(sources: list[dict[str, Any]]) -> tuple[list[dict], list[dict]]:
    escalations = []
    needs_confirmation = []
    for src in sources:
        if src.get("needs_manual_review"):
            escalations.append({
                "type": "scanned_pdf_or_unreadable",
                "detail": "텍스트 레이어가 없는 페이지가 있어 자동 추출 불가 (OCR 임의 수행 안 함)",
                "source": src["file"],
            })
        for chart in src.get("charts_detected", []):
            escalations.append({
                "type": "chart_source_unavailable",
                "detail": chart["note"],
                "source": chart["location"],
            })
        if src.get("structure_signal") == "none_detected" and src.get("type") in ("docx", "pdf"):
            needs_confirmation.append({
                "detail": "이 파일에서 제목/소제목 등 구조 신호를 찾지 못해 Content Group 계층을 만들지 못했습니다. "
                           "이후 [3] 단계에서 슬라이드 분할·주제 그룹핑을 수동 판단 근거 없이 진행하게 됩니다.",
                "source": src["file"],
            })
    return escalations, needs_confirmation


def main() -> None:
    parser = argparse.ArgumentParser(description="원본 자료에서 텍스트/표/이미지를 추출한다.")
    parser.add_argument("--input", nargs="+", required=True, help="원본 파일 또는 폴더 경로(복수 가능)")
    parser.add_argument("--output", required=True, help="출력 material_analysis.json 경로")
    args = parser.parse_args()

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    image_dir = output_path.parent / "extracted_images"
    image_dir.mkdir(parents=True, exist_ok=True)

    files = collect_files(args.input)
    if not files:
        raise SystemExit("처리할 파일이 없습니다.")

    sources = []
    for f in files:
        ext = f.suffix.lower()
        if ext in IMAGE_EXTS:
            sources.append(extract_image(f, image_dir))
        elif ext in EXTRACTORS:
            sources.append(EXTRACTORS[ext](f, image_dir))
        else:
            raise ValueError(f"지원하지 않는 파일 형식: {f}")

    escalations, needs_confirmation = build_escalations_and_confirmations(sources)

    result = {
        "sources": sources,
        "escalations": escalations,
        "needs_confirmation": needs_confirmation,
    }

    output_path.write_text(json.dumps(result, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
    print(f"[material-analysis] {len(files)}개 파일 처리 완료 -> {output_path}")
    if escalations:
        print(f"[material-analysis] 에스컬레이션 {len(escalations)}건 발생 — 사용자 확인 필요")


if __name__ == "__main__":
    main()
