"""
material-analysis / extract.py

원본 자료(docx/pdf/xlsx/csv/이미지)에서 텍스트·표·이미지를 추출해
material_analysis.json 스키마의 원재료(sources)를 생성한다.

원칙(설계서 1.6 준수):
- 수치는 원본 그대로 추출한다. 재계산·추정하지 않는다.
- 스캔 PDF(텍스트 레이어 없음) 등 자동 판독이 불가능한 경우 OCR을 임의로
  시도하지 않고 needs_manual_review 플래그만 남긴다.
- 임베드 차트는 원본 수치를 추출하지 않고 "원본 활용 필요"로만 표시한다
  (엑셀 차트의 내부 계열 데이터를 임의로 읽어 재구성하지 않음 — 실수로
  잘못된 값을 만들어낼 위험을 피하기 위함. 필요 시 openpyxl의 차트 객체를
  통해 명시적으로 계열 데이터를 읽는 기능은 향후 확장 가능).

사용법:
    python extract.py --input <파일 또는 폴더> [<파일2> ...] --output <json 경로>
"""
from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

SUPPORTED_EXTS = {".docx", ".pdf", ".xlsx", ".csv", ".png", ".jpg", ".jpeg", ".gif", ".bmp"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".bmp"}


def extract_docx(path: Path, image_dir: Path) -> dict[str, Any]:
    from docx import Document

    doc = Document(str(path))
    text_blocks = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_blocks.append({"text": para.text, "style": para.style.name if para.style else None})

    tables = []
    for table in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        headers = rows[0] if rows else []
        body = rows[1:] if len(rows) > 1 else []
        tables.append({
            "headers": headers,
            "rows": body,
            "row_count": len(body),
            "needs_appendix": len(body) >= 30,
        })

    images = _extract_docx_images(doc, path, image_dir)

    return {
        "file": str(path),
        "type": "docx",
        "text_blocks": text_blocks,
        "tables": tables,
        "images": images,
        "charts_detected": [],
        "needs_manual_review": False,
    }


def _extract_docx_images(doc, src_path: Path, image_dir: Path) -> list[str]:
    saved = []
    rels = doc.part.rels
    idx = 0
    for rel in rels.values():
        if "image" in rel.reltype:
            idx += 1
            image_part = rel.target_part
            ext = image_part.content_type.split("/")[-1]
            out_path = image_dir / f"{src_path.stem}_img{idx}.{ext}"
            out_path.write_bytes(image_part.blob)
            saved.append(str(out_path))
    return saved


def extract_pdf(path: Path, image_dir: Path) -> dict[str, Any]:
    import pdfplumber

    text_blocks = []
    tables = []
    images = []
    needs_manual_review = False

    with pdfplumber.open(str(path)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_blocks.append({"text": page_text, "page": page_num})
            else:
                # 텍스트 레이어가 없는 페이지 = 스캔 이미지일 가능성.
                # OCR을 임의로 수행하지 않고 수동 검토 대상으로만 표시한다.
                needs_manual_review = True

            for table in page.extract_tables():
                if not table:
                    continue
                headers, *body = table
                tables.append({
                    "headers": headers,
                    "rows": body,
                    "row_count": len(body),
                    "needs_appendix": len(body) >= 30,
                    "page": page_num,
                })

            for img_idx, img in enumerate(page.images, start=1):
                try:
                    cropped = page.within_bbox((img["x0"], img["top"], img["x1"], img["bottom"]))
                    im = cropped.to_image(resolution=200)
                    out_path = image_dir / f"{path.stem}_p{page_num}_img{img_idx}.png"
                    im.save(str(out_path))
                    images.append(str(out_path))
                except Exception:
                    # 페이지 이미지 추출 실패는 치명적이지 않음 — 건너뛰고 계속 진행
                    continue

    return {
        "file": str(path),
        "type": "pdf",
        "text_blocks": text_blocks,
        "tables": tables,
        "images": images,
        "charts_detected": [],
        "needs_manual_review": needs_manual_review,
    }


def extract_xlsx(path: Path, image_dir: Path) -> dict[str, Any]:
    import openpyxl

    wb = openpyxl.load_workbook(str(path), data_only=True)
    tables = []
    charts_detected = []

    for sheet in wb.worksheets:
        rows = [[cell for cell in row] for row in sheet.iter_rows(values_only=True)]
        rows = [r for r in rows if any(v is not None for v in r)]
        if rows:
            headers = list(rows[0])
            body = [list(r) for r in rows[1:]]
            tables.append({
                "sheet": sheet.title,
                "headers": headers,
                "rows": body,
                "row_count": len(body),
                "needs_appendix": len(body) >= 30,
            })

        # 차트 객체가 있으면 원본 수치를 임의로 재구성하지 않고 표시만 한다.
        if getattr(sheet, "_charts", None):
            for _ in sheet._charts:
                charts_detected.append({
                    "location": f"{path.name}!{sheet.title}",
                    "note": "차트 감지됨 — 원본 수치 추출 불가. 원본 파일 활용 또는 사용자에게 원본 데이터 확인 필요",
                })

    return {
        "file": str(path),
        "type": "xlsx",
        "text_blocks": [],
        "tables": tables,
        "images": [],
        "charts_detected": charts_detected,
        "needs_manual_review": False,
    }


def extract_csv(path: Path, image_dir: Path) -> dict[str, Any]:
    import csv

    with open(path, newline="", encoding="utf-8-sig") as f:
        reader = list(csv.reader(f))

    headers = reader[0] if reader else []
    body = reader[1:] if len(reader) > 1 else []

    return {
        "file": str(path),
        "type": "csv",
        "text_blocks": [],
        "tables": [{
            "headers": headers,
            "rows": body,
            "row_count": len(body),
            "needs_appendix": len(body) >= 30,
        }],
        "images": [],
        "charts_detected": [],
        "needs_manual_review": False,
    }


def extract_image(path: Path, image_dir: Path) -> dict[str, Any]:
    # 이미지 자체의 내용 판단(무엇을 담고 있는지)은 LLM이 이미지를 직접 보고 수행한다.
    # 여기서는 경로만 등록한다.
    return {
        "file": str(path),
        "type": "image",
        "text_blocks": [],
        "tables": [],
        "images": [str(path)],
        "charts_detected": [],
        "needs_manual_review": False,
    }


EXTRACTORS = {
    ".docx": extract_docx,
    ".pdf": extract_pdf,
    ".xlsx": extract_xlsx,
    ".csv": extract_csv,
}


def collect_files(inputs: list[str]) -> list[Path]:
    files: list[Path] = []
    for raw in inputs:
        p = Path(raw)
        if p.is_dir():
            files.extend(sorted(x for x in p.rglob("*") if x.suffix.lower() in SUPPORTED_EXTS))
        elif p.is_file():
            files.append(p)
        else:
            raise FileNotFoundError(f"입력 경로를 찾을 수 없습니다: {raw}")
    return files


def build_escalations_and_confirmations(sources: list[dict[str, Any]]) -> tuple[list[dict], list[dict]]:
    escalations = []
    needs_confirmation = []
    for src in sources:
        if src.get("needs_manual_review"):
            escalations.append({
                "type": "scanned_pdf_or_unreadable",
                "detail": "텍스트 레이어가 없는 페이지가 있어 자동 추출 불가 (OCR 임의 수행 안 함)",
                "source": src["file"],
            })
        for chart in src.get("charts_detected", []):
            escalations.append({
                "type": "chart_source_unavailable",
                "detail": chart["note"],
                "source": chart["location"],
            })
    return escalations, needs_confirmation


def main() -> None:
    parser = argparse.ArgumentParser(description="원본 자료에서 텍스트/표/이미지를 추출한다.")
    parser.add_argument("--input", nargs="+", required=True, help="원본 파일 또는 폴더 경로(복수 가능)")
    parser.add_argument("--output", required=True, help="출력 material_analysis.json 경로")
    args = parser.parse_args()

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    image_dir = output_path.parent / "extracted_images"
    image_dir.mkdir(parents=True, exist_ok=True)

    files = collect_files(args.input)
    if not files:
        raise SystemExit("처리할 파일이 없습니다.")

    sources = []
    for f in files:
        ext = f.suffix.lower()
        if ext in IMAGE_EXTS:
            sources.append(extract_image(f, image_dir))
        elif ext in EXTRACTORS:
            sources.append(EXTRACTORS[ext](f, image_dir))
        else:
            raise ValueError(f"지원하지 않는 파일 형식: {f}")

    escalations, needs_confirmation = build_escalations_and_confirmations(sources)

    result = {
        "sources": sources,
        "escalations": escalations,
        "needs_confirmation": needs_confirmation,
    }

    output_path.write_text(json.dumps(result, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
    print(f"[material-analysis] {len(files)}개 파일 처리 완료 -> {output_path}")
    if escalations:
        print(f"[material-analysis] 에스컬레이션 {len(escalations)}건 발생 — 사용자 확인 필요")


if __name__ == "__main__":
    main()
